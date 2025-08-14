"""
File Handler Utilities
=======================
Professional file handling utilities for EchoVerse.
Supports multiple file formats and text extraction.
"""

import os
import logging
from typing import List, Optional
import mimetypes
from pathlib import Path

logger = logging.getLogger(__name__)

class FileHandler:
    """
    Professional File Handler for EchoVerse
    
    Features:
    - Multiple file format support
    - Text extraction from various formats
    - File validation and security
    - Metadata extraction
    - File size and type checking
    """
    
    def __init__(self):
        """Initialize the file handler"""
        self.allowed_extensions = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.rtf': 'application/rtf',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.pdf': 'application/pdf',
            '.html': 'text/html',
            '.htm': 'text/html'
        }
        
        self.max_file_size = 16 * 1024 * 1024  # 16MB
        
        # Try to import optional dependencies
        self.docx_available = self._check_docx_support()
        self.pdf_available = self._check_pdf_support()
        
        logger.info("File Handler initialized")
    
    def _check_docx_support(self) -> bool:
        """Check if python-docx is available"""
        try:
            import docx
            return True
        except ImportError:
            logger.warning("python-docx not available - .docx files not supported")
            return False
    
    def _check_pdf_support(self) -> bool:
        """Check if PyPDF2 is available"""
        try:
            import PyPDF2
            return True
        except ImportError:
            logger.warning("PyPDF2 not available - .pdf files not supported")
            return False
    
    def allowed_file(self, filename: str) -> bool:
        """
        Check if file type is allowed
        
        Args:
            filename: Name of the file
            
        Returns:
            True if file type is allowed
        """
        if not filename:
            return False
        
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.allowed_extensions
    
    def validate_file(self, filepath: str) -> dict:
        """
        Validate uploaded file
        
        Args:
            filepath: Path to the file
            
        Returns:
            Validation result dictionary
        """
        if not os.path.exists(filepath):
            return {
                'valid': False,
                'error': 'File does not exist'
            }
        
        # Check file size
        file_size = os.path.getsize(filepath)
        if file_size > self.max_file_size:
            return {
                'valid': False,
                'error': f'File too large. Maximum size: {self.max_file_size / (1024*1024):.1f}MB'
            }
        
        # Check file extension
        filename = os.path.basename(filepath)
        if not self.allowed_file(filename):
            return {
                'valid': False,
                'error': f'File type not supported. Allowed types: {", ".join(self.allowed_extensions.keys())}'
            }
        
        # Check MIME type
        mime_type, _ = mimetypes.guess_type(filepath)
        file_ext = Path(filepath).suffix.lower()
        expected_mime = self.allowed_extensions.get(file_ext)
        
        return {
            'valid': True,
            'file_size': file_size,
            'mime_type': mime_type,
            'extension': file_ext,
            'expected_mime': expected_mime
        }
    
    def extract_text(self, filepath: str) -> str:
        """
        Extract text from various file formats
        
        Args:
            filepath: Path to the file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(filepath):
            logger.error(f"File not found: {filepath}")
            return ""
        
        # Validate file first
        validation = self.validate_file(filepath)
        if not validation['valid']:
            logger.error(f"File validation failed: {validation['error']}")
            return ""
        
        file_ext = validation['extension']
        
        try:
            if file_ext == '.txt':
                return self._extract_from_txt(filepath)
            elif file_ext == '.md':
                return self._extract_from_markdown(filepath)
            elif file_ext == '.rtf':
                return self._extract_from_rtf(filepath)
            elif file_ext == '.docx':
                return self._extract_from_docx(filepath)
            elif file_ext == '.pdf':
                return self._extract_from_pdf(filepath)
            elif file_ext in ['.html', '.htm']:
                return self._extract_from_html(filepath)
            else:
                logger.error(f"Unsupported file type: {file_ext}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return ""
    
    def _extract_from_txt(self, filepath: str) -> str:
        """Extract text from plain text file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading text file: {str(e)}")
                break
        
        return ""
    
    def _extract_from_markdown(self, filepath: str) -> str:
        """Extract text from Markdown file"""
        try:
            # For now, treat as plain text
            # Could be enhanced to parse Markdown and extract clean text
            text = self._extract_from_txt(filepath)
            
            # Basic Markdown cleanup
            import re
            
            # Remove headers
            text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
            
            # Remove bold/italic markers
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Italic
            text = re.sub(r'__([^_]+)__', r'\1', text)      # Bold
            text = re.sub(r'_([^_]+)_', r'\1', text)        # Italic
            
            # Remove links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            
            # Remove code blocks
            text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting from Markdown: {str(e)}")
            return ""
    
    def _extract_from_rtf(self, filepath: str) -> str:
        """Extract text from RTF file (basic implementation)"""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Basic RTF parsing - remove RTF control codes
            import re
            
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            
            # Remove braces
            content = re.sub(r'[{}]', '', content)
            
            # Clean up whitespace
            content = re.sub(r'\s+', ' ', content)
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from RTF: {str(e)}")
            return ""
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Extract text from DOCX file"""
        if not self.docx_available:
            logger.error("python-docx not available for DOCX extraction")
            return ""
        
        try:
            import docx
            
            doc = docx.Document(filepath)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, filepath: str) -> str:
        """Extract text from PDF file"""
        if not self.pdf_available:
            logger.error("PyPDF2 not available for PDF extraction")
            return ""
        
        try:
            import PyPDF2
            
            text_content = []
            
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            return ""
    
    def _extract_from_html(self, filepath: str) -> str:
        """Extract text from HTML file"""
        try:
            from html.parser import HTMLParser
            
            class HTMLTextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text_content = []
                    self.skip_tags = {'script', 'style', 'head', 'title', 'meta'}
                    self.current_tag = None
                
                def handle_starttag(self, tag, attrs):
                    self.current_tag = tag
                
                def handle_endtag(self, tag):
                    self.current_tag = None
                
                def handle_data(self, data):
                    if self.current_tag not in self.skip_tags:
                        text = data.strip()
                        if text:
                            self.text_content.append(text)
                
                def get_text(self):
                    return ' '.join(self.text_content)
            
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            parser = HTMLTextExtractor()
            parser.feed(html_content)
            
            return parser.get_text()
            
        except Exception as e:
            logger.error(f"Error extracting from HTML: {str(e)}")
            return ""
    
    def get_file_info(self, filepath: str) -> dict:
        """
        Get comprehensive file information
        
        Args:
            filepath: Path to the file
            
        Returns:
            File information dictionary
        """
        if not os.path.exists(filepath):
            return {'error': 'File not found'}
        
        try:
            stat = os.stat(filepath)
            filename = os.path.basename(filepath)
            file_ext = Path(filepath).suffix.lower()
            
            # Get MIME type
            mime_type, encoding = mimetypes.guess_type(filepath)
            
            # Extract text to get word count
            text_content = self.extract_text(filepath)
            word_count = len(text_content.split()) if text_content else 0
            
            return {
                'filename': filename,
                'extension': file_ext,
                'size_bytes': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'mime_type': mime_type,
                'encoding': encoding,
                'word_count': word_count,
                'character_count': len(text_content),
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime,
                'is_supported': self.allowed_file(filename)
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {'error': str(e)}
    
    def get_supported_formats(self) -> dict:
        """
        Get information about supported file formats
        
        Returns:
            Dictionary of supported formats and their capabilities
        """
        formats = {}
        
        for ext, mime_type in self.allowed_extensions.items():
            capabilities = ['text_extraction']
            
            if ext == '.docx' and self.docx_available:
                capabilities.append('advanced_formatting')
            elif ext == '.pdf' and self.pdf_available:
                capabilities.append('multi_page')
            elif ext == '.md':
                capabilities.append('markdown_parsing')
            elif ext in ['.html', '.htm']:
                capabilities.append('html_parsing')
            
            formats[ext] = {
                'mime_type': mime_type,
                'capabilities': capabilities,
                'available': True
            }
        
        # Mark unavailable formats
        if not self.docx_available:
            if '.docx' in formats:
                formats['.docx']['available'] = False
                formats['.docx']['note'] = 'Requires python-docx package'
        
        if not self.pdf_available:
            if '.pdf' in formats:
                formats['.pdf']['available'] = False
                formats['.pdf']['note'] = 'Requires PyPDF2 package'
        
        return formats
    
    def cleanup_temp_files(self, directory: str, max_age_hours: int = 24):
        """
        Clean up temporary files older than specified age
        
        Args:
            directory: Directory to clean
            max_age_hours: Maximum age in hours
        """
        try:
            import time
            
            if not os.path.exists(directory):
                return
            
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for filename in os.listdir(directory):
                filepath = os.path.join(directory, filename)
                
                if os.path.isfile(filepath):
                    file_age = current_time - os.path.getmtime(filepath)
                    
                    if file_age > max_age_seconds:
                        try:
                            os.remove(filepath)
                            logger.info(f"Cleaned up old file: {filename}")
                        except Exception as e:
                            logger.warning(f"Failed to remove {filename}: {str(e)}")
            
        except Exception as e:
            logger.error(f"Error during cleanup: {str(e)}")
